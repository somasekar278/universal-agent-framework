"""
AI-Powered Architecture Advisor

Analyzes natural language use case briefs and recommends the optimal
SOTA Agent Framework architecture, including:
- Learning level (1-5)
- Schemas to use
- Features to enable
- Integrations needed
- Architecture patterns

Usage:
    # From text
    sota-architect "I need a fraud detection system with memory and self-improvement"
    
    # From file
    sota-architect --file requirements.txt
    sota-architect --file project_brief.pdf
    
    # Interactive
    sota-architect --interactive
"""

import re
import os
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
from enum import Enum


class DocumentParser:
    """
    Parse documents of various formats to extract text for analysis.
    
    Supports:
    - Plain text (.txt)
    - Markdown (.md)
    - PDF (.pdf) - requires PyPDF2
    - Word (.docx, .doc) - requires python-docx
    """
    
    @staticmethod
    def validate_file(file_path: str) -> tuple[bool, str]:
        """
        Validate file before parsing.
        
        Returns:
            (is_valid, error_message)
        """
        path = Path(file_path)
        
        # Check if file exists
        if not path.exists():
            return False, f"File not found: {file_path}"
        
        # Check if it's a file (not directory)
        if not path.is_file():
            return False, f"Path is not a file: {file_path}"
        
        # Check file size (warn if too large > 10MB)
        file_size = path.stat().st_size
        if file_size > 10 * 1024 * 1024:  # 10MB
            return False, (
                f"File is very large ({file_size / (1024*1024):.1f} MB). "
                "Consider using a smaller document or extracting key sections."
            )
        
        # Check if file is empty
        if file_size == 0:
            return False, "File is empty (0 bytes)"
        
        # Check file extension
        suffix = path.suffix.lower()
        supported = ['.txt', '.md', '.markdown', '.rst', '.pdf', '.docx', '.doc']
        if suffix not in supported:
            return False, (
                f"Unsupported file format: {suffix}\n"
                f"Supported: {', '.join(supported)}"
            )
        
        return True, ""
    
    @staticmethod
    def parse_file(file_path: str) -> str:
        """
        Parse a file and extract text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is unsupported or file is invalid
        """
        # Validate file first
        is_valid, error_msg = DocumentParser.validate_file(file_path)
        if not is_valid:
            if "not found" in error_msg.lower():
                raise FileNotFoundError(error_msg)
            else:
                raise ValueError(error_msg)
        
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        # Plain text and markdown
        if suffix in ['.txt', '.md', '.markdown', '.rst']:
            return DocumentParser._parse_text(path)
        
        # PDF
        elif suffix == '.pdf':
            return DocumentParser._parse_pdf(path)
        
        # Word documents
        elif suffix in ['.docx', '.doc']:
            return DocumentParser._parse_docx(path)
        
        else:
            raise ValueError(
                f"Unsupported file format: {suffix}\n"
                f"Supported: .txt, .md, .pdf, .docx, .doc"
            )
    
    @staticmethod
    def _parse_text(path: Path) -> str:
        """Parse plain text or markdown file."""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    content = f.read()
                    
                    # Check if file is empty
                    if not content.strip():
                        raise ValueError("File is empty or contains only whitespace.")
                    
                    return content
                    
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise RuntimeError(f"Error reading file: {str(e)}")
        
        # If all encodings failed
        raise ValueError(
            "Unable to decode file. It may be binary or use an unsupported encoding.\n"
            "Try converting to UTF-8 first."
        )
    
    @staticmethod
    def _parse_pdf(path: Path) -> str:
        """Parse PDF file (requires PyPDF2)."""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError(
                "PDF parsing requires PyPDF2.\n"
                "Install with: pip install PyPDF2\n"
                "Or: pip install sota-agent-framework[documents]"
            )
        
        try:
            text_parts = []
            
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                # Check if PDF is encrypted
                if reader.is_encrypted:
                    # Try empty password first
                    try:
                        reader.decrypt('')
                    except:
                        raise ValueError(
                            "PDF is password-protected. Please provide an unencrypted version."
                        )
                
                # Check if PDF has pages
                if len(reader.pages) == 0:
                    raise ValueError("PDF has no pages or is empty.")
                
                # Extract text from each page
                for page_num, page in enumerate(reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            text_parts.append(text)
                        else:
                            # Page might be image-only
                            text_parts.append(f"[Page {page_num}: No extractable text - possibly image-only]")
                    except Exception as page_error:
                        text_parts.append(f"[Page {page_num}: Error extracting text: {str(page_error)}]")
                
                # Check if we got any meaningful text
                full_text = '\n'.join(text_parts)
                if not full_text.strip() or full_text.count('[Page') == len(reader.pages):
                    raise ValueError(
                        "PDF appears to contain only images. OCR is required.\n"
                        "Consider converting to text first or using an OCR tool."
                    )
                
                return full_text
                
        except PyPDF2.errors.PdfReadError as e:
            raise ValueError(f"Invalid or corrupted PDF file: {str(e)}")
        except ValueError as e:
            # Re-raise ValueError with our custom messages
            raise
        except Exception as e:
            raise RuntimeError(f"Unexpected error parsing PDF: {str(e)}")
    
    @staticmethod
    def _parse_docx(path: Path) -> str:
        """Parse Word document (requires python-docx)."""
        try:
            import docx
        except ImportError:
            raise ImportError(
                "Word document parsing requires python-docx.\n"
                "Install with: pip install python-docx\n"
                "Or: pip install sota-agent-framework[documents]"
            )
        
        try:
            doc = docx.Document(path)
            
            # Extract paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            # Check if we got any text
            full_text = '\n'.join(text_parts)
            if not full_text.strip():
                raise ValueError("Word document appears to be empty or contains no extractable text.")
            
            return full_text
            
        except ValueError as e:
            # Re-raise ValueError with our custom messages
            raise
        except Exception as e:
            raise RuntimeError(f"Error parsing Word document: {str(e)}")


class ComplexityLevel(Enum):
    """Complexity levels mapping to learning levels."""
    SIMPLE = 1          # Simple chatbot
    CONTEXTUAL = 2      # Context-aware with memory
    PRODUCTION = 3      # Production API with monitoring
    ADVANCED = 4        # Complex workflows with planning
    EXPERT = 5          # Multi-agent with A2A


@dataclass
class ArchitectureRecommendation:
    """Complete architecture recommendation."""
    level: ComplexityLevel
    level_name: str
    confidence: float
    
    # Schemas
    input_schema: str
    output_schema: str
    
    # Core features
    features: List[str]
    
    # Optional integrations
    integrations: List[str]
    
    # Reasoning
    reasoning: str
    
    # Estimated effort
    estimated_hours: str
    
    # Code generation params
    generation_params: Dict[str, Any]


class ArchitectureAdvisor:
    """
    AI-powered architecture advisor that analyzes use case briefs
    and recommends optimal SOTA Agent Framework architecture.
    
    Uses intelligent text analysis that:
    - Focuses on actual requirements, not titles/buzzwords
    - Weights technical sections higher than marketing language
    - Analyzes implementation details, not aspirational descriptions
    """
    
    @staticmethod
    def print_capabilities_overview():
        """Print an overview of all framework levels and their capabilities."""
        print("\n" + "="*80)
        print("📚 SOTA Agent Framework - Complexity Levels Overview")
        print("="*80)
        print()
        print("┌─────────┬──────────────────────┬─────────────┬────────────────────────────────┐")
        print("│  Level  │   Architecture       │   Time      │  Key Capabilities              │")
        print("├─────────┼──────────────────────┼─────────────┼────────────────────────────────┤")
        print("│ Level 1 │ Simple Chatbot       │   2-4 hrs   │ • Basic Q&A                    │")
        print("│         │                      │             │ • FAQ responses                │")
        print("│         │                      │             │ • No memory                    │")
        print("├─────────┼──────────────────────┼─────────────┼────────────────────────────────┤")
        print("│ Level 2 │ Context-Aware        │   4-8 hrs   │ • Session memory               │")
        print("│         │                      │             │ • Personalization              │")
        print("│         │                      │             │ • History tracking             │")
        print("├─────────┼──────────────────────┼─────────────┼────────────────────────────────┤")
        print("│ Level 3 │ Production API       │  8-16 hrs   │ • High scale (1000s req/sec)   │")
        print("│         │                      │             │ • 99.9% uptime                 │")
        print("│         │                      │             │ • Full monitoring              │")
        print("├─────────┼──────────────────────┼─────────────┼────────────────────────────────┤")
        print("│ Level 4 │ Complex Workflow     │ 16-32 hrs   │ • Plan-Act-Critique loops      │")
        print("│         │                      │             │ • Self-improvement             │")
        print("│         │                      │             │ • Feedback optimization        │")
        print("├─────────┼──────────────────────┼─────────────┼────────────────────────────────┤")
        print("│ Level 5 │ Multi-Agent System   │ 32-64 hrs   │ • Multiple autonomous agents   │")
        print("│         │                      │             │ • Agent-to-agent communication │")
        print("│         │                      │             │ • Distributed coordination     │")
        print("└─────────┴──────────────────────┴─────────────┴────────────────────────────────┘")
        print()
        print("💡 Each level builds on previous capabilities")
        print("⚠️  Recommendations are guidance - choose what fits your needs")
        print("="*80 + "\n")
    
    def __init__(self):
        """Initialize the advisor with pattern matchers and rules."""
        
        # Complexity indicators (patterns that suggest higher complexity)
        # Focus on implementation requirements, not aspirational buzzwords
        self.complexity_patterns = {
            ComplexityLevel.SIMPLE: [
                r'\bsimple\b', r'\bbasic\b', r'\bquick\b', r'\bchatbot\b',
                r'\bfaq\b', r'\bquestion.?answer', r'\brespond\b',
                r'\bsingle\s+purpose\b', r'\bstraightforward\b'
            ],
            ComplexityLevel.CONTEXTUAL: [
                r'\bremember\s+.{0,30}(history|context|conversation|preferences|what|topics)',
                r'\btrack.{0,30}(progress|session|user|student|patient).{0,30}(across|over|throughout)',
                r'\bpersonaliz',
                r'\bconversation\s+history\b', r'\bmulti.?turn\b',
                r'\bstore\s+.{0,20}(data|information|preferences)',
                r'\bretain\s+.{0,20}information',
                r'\bcontext\s+(across|throughout)\s+.{0,20}(sessions|conversations|journey)',
                r'\bmaintain.{0,20}context.{0,20}throughout',
                r'\badapt.{0,30}based\s+on.{0,30}(their|user|history)'
            ],
            ComplexityLevel.PRODUCTION: [
                r'\bproduction.?grade\b', r'\brest\s+api\b', r'\bmicroservice\b',
                r'\bmonitor', r'\bscale\s+to\b', r'\bhealth.?check\b',
                r'\b99\.9+%\s+uptime\b', r'\bload\s+balanc\b',
                r'\bauthentication\b', r'\bauthorization\b'
            ],
            ComplexityLevel.ADVANCED: [
                r'\bplan\s+.{0,30}execut\w+.{0,30}(critiqu|evaluat|assess)',
                r'\bmulti.?step\s+.{0,20}(workflow|process|task)',
                r'\boptimiz\w+\s+.{0,30}(over\s+time|based\s+on|through)',
                r'\bself.?improv|continuously.?improv|automatically.?improv',
                r'\bfeedback\s+loop',
                r'\bcritiqu\w+.{0,30}(result|output|performance)',
                r'\badaptive\s+(behavior|strateg|approach)',
                r'\blearn\w+\s+from\s+.{0,30}(mistake|outcome|feedback|result|engagement)',
                r'\biterat\w+\s+(improve|refine|enhance|adjust)',
                r'\brevise.{0,20}strateg', r'\b(refine|adjust)\s+.{0,20}(approach|strategy|plan)',
                r'\bimprove.{0,20}strateg.{0,20}over\s+time',
                r'\bevaluat\w+.{0,20}(quality|performance).{0,20}(and|then).{0,20}(improv|adjust|refine)'
            ],
            ComplexityLevel.EXPERT: [
                r'\bmultiple\s+.{0,30}agents?\s+.{0,50}(communicate|collaborate|coordinate)',
                r'\bagents?\s+.{0,30}(communicate|collaborate|coordinate)\s+.{0,30}(with\s+)?(each\s+other|together)',
                r'\bagent.?to.?agent',
                r'\bdistributed\s+.{0,20}(system|agents?|network)', 
                r'\bpeer.?to.?peer',
                r'\bagent\s+discovery\b', 
                r'\bagents?\s+.{0,30}(discover|find)\s+.{0,30}(each\s+other|other\s+agents)',
                r'\bcross.?platform\s+.{0,30}agents?\b',
                r'\bdecentralized\s+.{0,30}(coordination|agents?)\b', 
                r'\bautonomous\s+agents?\s+.{0,30}(collaborate|coordinate|work\s+together)',
                r'\bagents?\s+.{0,30}request\s+.{0,20}(help|consultation)',
                r'\bseparate\s+agents?\s+for\b',
                r'\bspecialized\s+agents?\b.{0,50}(coordinate|collaborate|work\s+together)',
                r'\bagents?\s+for\s+(different|various)\s+.{0,30}(must|should|need).{0,50}(coordinate|collaborate|communicate)'
            ]
        }
        
        # Feature indicators - focus on actual implementation needs
        self.feature_patterns = {
            'memory': [
                r'\bstore\s+.{0,20}(history|context|data|progress)',
                r'\bremember\s+.{0,20}(previous|past|earlier|what|topic)',
                r'\bretain\s+.{0,20}information',
                r'\btrack\s+.{0,30}(session|user|conversation|progress|student|patient)',
                r'\bcontext\s+across\s+.{0,20}(sessions|interactions|conversation)',
                r'\bmaintain.{0,20}context',
                r'\btrack.{0,20}(weak\s+areas|preferences|history|journey)'
            ],
            'langgraph': [
                r'\bmulti.?step\s+\w+\s+(workflow|process|pipeline)\b',
                r'\bplan\s+and\s+execute\b', r'\bstate\s+machine\b',
                r'\borchestrat\w+\s+\w+\s+(tasks|steps)\b',
                r'\bsequential\s+\w+\s+processing\b',
                r'\bworkflow\s+\w+\s+(coordination|management)\b'
            ],
            'mcp': [
                r'\bexternal\s+\w+\s+(api|service|tool)\b',
                r'\bintegrat\w+\s+with\s+\w+\s+(third.?party|external)\b',
                r'\bcall\s+\w+\s+(external|remote)\b',
                r'\btool\s+\w+\s+interface\b'
            ],
            'a2a': [
                r'\bagents?\s+.{0,30}(communicate|collaborate|coordinate)\s+.{0,30}(with\s+)?each\s+other',
                r'\bpeer.?to.?peer',
                r'\bagent.?to.?agent',
                r'\bdistributed\s+.{0,20}agents?\b',
                r'\bagents?\s+.{0,30}discover\s+.{0,30}(each\s+other|other\s+agents)',
                r'\bcross.?framework\s+.{0,30}(agent|communication)',
                r'\bagents?\s+.{0,30}request\s+.{0,20}help'
            ],
            'monitoring': [
                r'\bmonitor\s+\w+\s+(performance|health|metrics)\b',
                r'\bobservability\s+\w+\s+(platform|system)\b',
                r'\btrace\s+\w+\s+(execution|requests)\b',
                r'\btrack\s+\w+\s+(metrics|errors|latency)\b',
                r'\btelemetry\s+\w+\s+data\b'
            ],
            'optimization': [
                r'\boptimiz\w+\s+.{0,30}(performance|prompts|responses|strategies|criteria)',
                r'\bimprov\w+\s+.{0,30}(over\s+time|through|based\s+on|its\s+\w+|strategies)',
                r'\btune\s+.{0,20}(parameters|models)',
                r'\bself.?learn\w+\s+from\s+.{0,20}(feedback|data)',
                r'\bfeedback\s+loop', 
                r'\badaptive\s+.{0,20}(behavior|strategies|approach)',
                r'\blearn\s+from\s+.{0,30}(accepted|rejected|performance)',
                r'\bcontinuously.{0,20}improve',
                r'\bautomatically.{0,20}adjust',
                r'\brefine.{0,20}(criteria|strategies|approach).{0,20}over\s+time'
            ],
            'benchmarking': [
                r'\bmeasure\s+\w+\s+(accuracy|performance|quality)\b',
                r'\beval\w+\s+\w+\s+(model|agent|system)\b',
                r'\btest\s+\w+\s+(quality|performance)\b',
                r'\bquality\s+\w+\s+(metrics|assessment)\b'
            ],
            'databricks': [
                r'\bdatabricks\s+\w+\s+(platform|integration|deployment)\b',
                r'\bunity\s+catalog\s+\w+\s+(for|integration)\b',
                r'\bmlflow\s+\w+\s+(tracking|registry)\b',
                r'\bspark\s+\w+\s+(processing|cluster)\b',
                r'\bdelta\s+lake\s+\w+\s+(storage|tables)\b'
            ]
        }
        
        # Domain indicators
        self.domain_patterns = {
            'fraud': [r'\bfraud\b', r'\bscam\b', r'\brisk\b', r'\bsuspicious\b'],
            'customer_support': [r'\bsupport\b', r'\bticket\b', r'\bhelp.?desk\b', r'\bcustomer\b'],
            'analytics': [r'\banalytics\b', r'\binsight\b', r'\bdata\b', r'\breport\b'],
            'healthcare': [r'\bhealth', r'\bmedical\b', r'\bpatient\b', r'\bdiagnos'],
            'finance': [r'\bfinance\b', r'\bbanking\b', r'\btrade\b', r'\binvest'],
            'ecommerce': [r'\becommerce\b', r'\bshopping\b', r'\bproduct\b', r'\bcart\b'],
            'hr': [r'\bhr\b', r'\brecruit', r'\bhiring\b', r'\bemployee\b'],
            'legal': [r'\blegal\b', r'\bcontract\b', r'\bcompliance\b', r'\bregulatr']
        }
    
    def analyze_brief(self, brief: str) -> ArchitectureRecommendation:
        """
        Analyze a use case brief and recommend architecture.
        
        Args:
            brief: Natural language description of the use case
            
        Returns:
            ArchitectureRecommendation with all details
        """
        # Step 0: Preprocess the brief to extract meaningful content
        processed_brief = self._preprocess_brief(brief)
        brief_lower = processed_brief.lower()
        
        # Step 1: Determine complexity level
        level, level_confidence = self._determine_complexity(brief_lower)
        
        # Step 2: Identify required features
        features = self._identify_features(brief_lower, level)
        
        # Step 3: Identify integrations
        integrations = self._identify_integrations(brief_lower, level)
        
        # Step 4: Recommend schemas
        input_schema, output_schema = self._recommend_schemas(level)
        
        # Step 5: Detect domain
        domain = self._detect_domain(brief_lower)
        
        # Step 6: Generate reasoning
        reasoning = self._generate_reasoning(brief, level, features, integrations, domain)
        
        # Step 7: Estimate effort
        estimated_hours = self._estimate_effort(level, features)
        
        # Step 8: Create generation params
        generation_params = self._create_generation_params(
            level, domain, features, integrations
        )
        
        return ArchitectureRecommendation(
            level=level,
            level_name=self._get_level_name(level),
            confidence=level_confidence,
            input_schema=input_schema,
            output_schema=output_schema,
            features=features,
            integrations=integrations,
            reasoning=reasoning,
            estimated_hours=estimated_hours,
            generation_params=generation_params
        )
    
    def _preprocess_brief(self, brief: str) -> str:
        """
        Preprocess brief to focus on actual requirements, not titles/buzzwords.
        
        Strategy:
        - Extract sections with actual requirements (requirements, technical, implementation)
        - De-weight or remove titles, headers, executive summaries
        - Focus on concrete descriptions of what needs to be built
        - Remove aspirational/marketing language
        """
        lines = brief.split('\n')
        processed_lines = []
        
        # Markers for high-value sections (actual requirements)
        high_value_markers = [
            'requirement', 'technical', 'implementation', 'functionality',
            'feature', 'must', 'should', 'need', 'will', 'capability',
            'specification', 'architecture', 'design', 'component'
        ]
        
        # Markers for low-value sections (titles, marketing)
        low_value_markers = [
            'executive summary', 'overview', 'introduction', 'background',
            'vision', 'mission', 'goal', 'objective', 'title', 'brief'
        ]
        
        in_high_value_section = False
        in_low_value_section = False
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Skip empty lines
            if not line_lower:
                continue
            
            # Skip very short lines (likely titles/headers)
            if len(line_lower) < 20 and ':' not in line_lower:
                # But check if it's a section header
                if any(marker in line_lower for marker in high_value_markers):
                    in_high_value_section = True
                    in_low_value_section = False
                    continue
                elif any(marker in line_lower for marker in low_value_markers):
                    in_low_value_section = True
                    in_high_value_section = False
                    continue
                else:
                    # Skip standalone short lines (titles)
                    continue
            
            # Skip low-value sections
            if in_low_value_section and len(line_lower) < 100:
                # Short lines in low-value sections are skipped
                continue
            
            # Check if line contains high-value content
            is_high_value = any(marker in line_lower for marker in high_value_markers)
            is_low_value = any(marker in line_lower for marker in low_value_markers)
            
            if is_high_value or in_high_value_section:
                # Include this line with full weight
                processed_lines.append(line)
            elif not is_low_value:
                # Neutral line - include but maybe de-weight later
                processed_lines.append(line)
        
        # If we didn't find any high-value sections, use original brief
        # (might be a short, direct description)
        if len(processed_lines) < 5:
            return brief
        
        return '\n'.join(processed_lines)
    
    def _determine_complexity(self, brief: str) -> Tuple[ComplexityLevel, float]:
        """Determine complexity level from brief."""
        scores = {}
        
        # Score each level based on pattern matches
        for level, patterns in self.complexity_patterns.items():
            score = sum(1 for pattern in patterns if re.search(pattern, brief))
            scores[level] = score
        
        # Get level with highest score
        if not any(scores.values()):
            # Default to SIMPLE if no patterns match
            return ComplexityLevel.SIMPLE, 0.5
        
        max_level = max(scores.items(), key=lambda x: x[1])
        
        # Calculate confidence (0-1)
        total_matches = sum(scores.values())
        confidence = min(max_level[1] / max(total_matches, 1), 1.0)
        
        # Boost confidence if multiple levels suggest same direction
        if max_level[1] > 0:
            confidence = min(confidence + 0.2, 1.0)
        
        return max_level[0], confidence
    
    def _identify_features(self, brief: str, level: ComplexityLevel) -> List[str]:
        """Identify required features from brief."""
        features = []
        
        for feature, patterns in self.feature_patterns.items():
            if any(re.search(pattern, brief) for pattern in patterns):
                features.append(feature)
        
        # Add default features based on level
        level_defaults = {
            ComplexityLevel.SIMPLE: [],
            ComplexityLevel.CONTEXTUAL: ['memory'],
            ComplexityLevel.PRODUCTION: ['memory', 'monitoring'],
            ComplexityLevel.ADVANCED: ['memory', 'monitoring', 'langgraph', 'optimization'],
            ComplexityLevel.EXPERT: ['memory', 'monitoring', 'langgraph', 'optimization', 'benchmarking']
        }
        
        # Add defaults if not already present
        for default in level_defaults.get(level, []):
            if default not in features:
                features.append(default)
        
        return features
    
    def _identify_integrations(self, brief: str, level: ComplexityLevel) -> List[str]:
        """Identify required integrations."""
        integrations = []
        
        # Check for explicit integration mentions
        if any(re.search(pattern, brief) for pattern in self.feature_patterns['mcp']):
            integrations.append('MCP')
        
        if any(re.search(pattern, brief) for pattern in self.feature_patterns['a2a']):
            integrations.append('A2A')
        
        if any(re.search(pattern, brief) for pattern in self.feature_patterns['databricks']):
            integrations.append('Databricks')
        
        # Add LangGraph for Advanced+ levels
        if level.value >= ComplexityLevel.ADVANCED.value:
            if 'LangGraph' not in integrations:
                integrations.append('LangGraph')
        
        return integrations
    
    def _recommend_schemas(self, level: ComplexityLevel) -> Tuple[str, str]:
        """Recommend input/output schemas based on level."""
        schema_map = {
            ComplexityLevel.SIMPLE: ('ChatInput', 'ChatOutput'),
            ComplexityLevel.CONTEXTUAL: ('ContextAwareInput', 'ContextAwareOutput'),
            ComplexityLevel.PRODUCTION: ('APIRequest', 'APIResponse'),
            ComplexityLevel.ADVANCED: ('WorkflowInput', 'WorkflowOutput'),
            ComplexityLevel.EXPERT: ('CollaborationRequest', 'CollaborationResponse')
        }
        return schema_map[level]
    
    def _detect_domain(self, brief: str) -> Optional[str]:
        """Detect the domain from the brief."""
        domain_scores = {}
        
        for domain, patterns in self.domain_patterns.items():
            score = sum(1 for pattern in patterns if re.search(pattern, brief))
            if score > 0:
                domain_scores[domain] = score
        
        if domain_scores:
            return max(domain_scores.items(), key=lambda x: x[1])[0]
        
        return 'general'
    
    def _generate_reasoning(
        self,
        brief: str,
        level: ComplexityLevel,
        features: List[str],
        integrations: List[str],
        domain: str
    ) -> str:
        """Generate human-readable reasoning for the recommendation."""
        reasoning_parts = []
        
        # Level reasoning
        level_reasons = {
            ComplexityLevel.SIMPLE: "Based on your requirements, a simple chatbot architecture is sufficient.",
            ComplexityLevel.CONTEXTUAL: "Your use case requires context awareness and memory capabilities.",
            ComplexityLevel.PRODUCTION: "This is a production-grade system requiring robust API design and monitoring.",
            ComplexityLevel.ADVANCED: "This is a complex workflow requiring planning, execution, and self-improvement loops.",
            ComplexityLevel.EXPERT: "This is an expert-level multi-agent system requiring advanced collaboration."
        }
        reasoning_parts.append(level_reasons[level])
        
        # Domain reasoning
        if domain != 'general':
            reasoning_parts.append(f"Detected domain: {domain.replace('_', ' ').title()}")
        
        # Feature reasoning
        if features:
            feature_names = ', '.join(features)
            reasoning_parts.append(f"Required features: {feature_names}")
        
        # Integration reasoning
        if integrations:
            integration_names = ', '.join(integrations)
            reasoning_parts.append(f"Recommended integrations: {integration_names}")
        
        return ' '.join(reasoning_parts)
    
    def _estimate_effort(self, level: ComplexityLevel, features: List[str]) -> str:
        """Estimate development effort."""
        base_hours = {
            ComplexityLevel.SIMPLE: (2, 4),
            ComplexityLevel.CONTEXTUAL: (4, 8),
            ComplexityLevel.PRODUCTION: (8, 16),
            ComplexityLevel.ADVANCED: (16, 32),
            ComplexityLevel.EXPERT: (32, 64)
        }
        
        min_hours, max_hours = base_hours[level]
        
        # Add hours for additional features
        extra_hours = len(features) * 2
        max_hours += extra_hours
        
        return f"{min_hours}-{max_hours} hours"
    
    def _create_generation_params(
        self,
        level: ComplexityLevel,
        domain: str,
        features: List[str],
        integrations: List[str]
    ) -> Dict[str, Any]:
        """Create parameters for code generation."""
        return {
            'level': level.value,
            'domain': domain,
            'features': {
                'memory': 'memory' in features,
                'langgraph': 'langgraph' in features,
                'mcp': 'mcp' in features or 'MCP' in integrations,
                'a2a': 'a2a' in features or 'A2A' in integrations,
                'monitoring': 'monitoring' in features,
                'optimization': 'optimization' in features,
                'benchmarking': 'benchmarking' in features,
                'databricks': 'databricks' in features or 'Databricks' in integrations
            },
            'schemas': {
                'input': self._recommend_schemas(level)[0],
                'output': self._recommend_schemas(level)[1]
            }
        }
    
    def _get_level_name(self, level: ComplexityLevel) -> str:
        """Get human-readable level name."""
        names = {
            ComplexityLevel.SIMPLE: "Level 1: Simple Chatbot",
            ComplexityLevel.CONTEXTUAL: "Level 2: Context-Aware Agent",
            ComplexityLevel.PRODUCTION: "Level 3: Production API",
            ComplexityLevel.ADVANCED: "Level 4: Complex Workflow",
            ComplexityLevel.EXPERT: "Level 5: Multi-Agent System"
        }
        return names[level]
    
    def prompt_level_selection(self, recommendation: ArchitectureRecommendation) -> int:
        """
        Interactive prompt for level selection with override capability.
        
        Returns:
            Selected level (1-5)
        """
        print("\n" + "="*70)
        print("🎯 Level Selection")
        print("="*70)
        print("\nAll available levels:")
        print()
        print("  1️⃣  Level 1: Simple Chatbot")
        print("      ↳ Basic Q&A, FAQ bots, simple responses")
        print("      ↳ Time: 2-4 hours | Features: None required")
        print()
        print("  2️⃣  Level 2: Context-Aware Agent")
        print("      ↳ Memory, personalization, session tracking")
        print("      ↳ Time: 4-8 hours | Features: Memory")
        print()
        print("  3️⃣  Level 3: Production API")
        print("      ↳ High scale, monitoring, 24/7 operations")
        print("      ↳ Time: 8-16 hours | Features: Memory, Monitoring")
        print()
        print("  4️⃣  Level 4: Complex Workflow")
        print("      ↳ Plan-Act-Critique, self-improvement, optimization")
        print("      ↳ Time: 16-32 hours | Features: LangGraph, Optimization, Memory")
        print()
        print("  5️⃣  Level 5: Multi-Agent System")
        print("      ↳ Multiple agents, A2A communication, distributed")
        print("      ↳ Time: 32-64 hours | Features: A2A, All advanced features")
        print()
        print("="*70)
        print(f"\n💡 Recommended: Level {recommendation.level.value} (confidence: {recommendation.confidence:.0%})")
        print()
        
        while True:
            try:
                choice = input("Select level (1-5) or press Enter for recommended: ").strip()
                
                if choice == "":
                    # Use recommendation
                    selected = recommendation.level.value
                    print(f"\n✅ Using recommended Level {selected}")
                    return selected
                
                selected = int(choice)
                if 1 <= selected <= 5:
                    if selected != recommendation.level.value:
                        print(f"\n⚠️  You selected Level {selected}, different from recommendation (Level {recommendation.level.value})")
                        confirm = input("   Proceed with Level " + str(selected) + "? (y/n): ").strip().lower()
                        if confirm in ['y', 'yes']:
                            print(f"\n✅ Using Level {selected}")
                            return selected
                        else:
                            print("\n   Let's choose again...")
                            continue
                    else:
                        print(f"\n✅ Using Level {selected}")
                        return selected
                else:
                    print("   ❌ Please enter a number between 1 and 5")
            except ValueError:
                print("   ❌ Please enter a valid number (1-5) or press Enter")
            except (KeyboardInterrupt, EOFError):
                print(f"\n\n✅ Using recommended Level {recommendation.level.value}")
                return recommendation.level.value
    
    def print_recommendation(self, rec: ArchitectureRecommendation, show_json: bool = True):
        """Pretty print the recommendation."""
        print("\n" + "="*70)
        print("🏗️  SOTA Agent Framework - Architecture Recommendation")
        print("="*70)
        
        print(f"\n📊 Recommended Level: {rec.level_name}")
        print(f"   Confidence: {rec.confidence:.0%}")
        print(f"   ⚠️  This is guidance only - you can choose any level that fits your needs")
        
        print(f"\n📋 Schemas:")
        print(f"   Input:  {rec.input_schema}")
        print(f"   Output: {rec.output_schema}")
        
        if rec.features:
            print(f"\n✨ Core Features:")
            for feature in rec.features:
                print(f"   • {feature}")
        
        if rec.integrations:
            print(f"\n🔌 Integrations:")
            for integration in rec.integrations:
                print(f"   • {integration}")
        
        print(f"\n💡 Reasoning:")
        print(f"   {rec.reasoning}")
        
        print(f"\n⏱️  Estimated Effort: {rec.estimated_hours}")
        
        print("\n" + "="*70)
        
        # Show detailed JSON output by default
        if show_json:
            print("\n📊 Detailed Analysis (JSON):")
            print("="*70)
            import json
            output = {
                'level': rec.level.value,
                'level_name': rec.level_name,
                'confidence': rec.confidence,
                'input_schema': rec.input_schema,
                'output_schema': rec.output_schema,
                'features': rec.features,
                'integrations': rec.integrations,
                'reasoning': rec.reasoning,
                'estimated_hours': rec.estimated_hours,
                'generation_params': rec.generation_params
            }
            print(json.dumps(output, indent=2))
            print("="*70)
        
        print()


def main():
    """CLI entry point for architecture advisor."""
    import sys
    import argparse
    
    parser = argparse.ArgumentParser(
        description='SOTA Agent Framework Architecture Advisor',
        epilog='Examples:\n'
               '  sota-architect "Build a fraud detection system"\n'
               '  sota-architect --file requirements.txt\n'
               '  sota-architect --file project_brief.pdf',
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument(
        'brief',
        nargs='*',
        help='Natural language description of your use case'
    )
    parser.add_argument(
        '-f', '--file',
        type=str,
        help='Path to document file (.txt, .md, .pdf, .docx, .doc)'
    )
    parser.add_argument(
        '-i', '--interactive',
        action='store_true',
        help='Interactive mode with prompts for input'
    )
    parser.add_argument(
        '--select',
        action='store_true',
        help='Interactive level selection after recommendation'
    )
    parser.add_argument(
        '-j', '--json',
        action='store_true',
        help='Output as JSON (non-interactive)'
    )
    
    args = parser.parse_args()
    
    advisor = ArchitectureAdvisor()
    
    # Get brief
    if args.file:
        # Parse document file
        print(f"\n📄 Parsing document: {args.file}")
        try:
            brief = DocumentParser.parse_file(args.file)
            
            print(f"✅ Extracted {len(brief)} characters")
            
            # Show preview if not JSON output
            if not args.json:
                preview = brief[:200] + "..." if len(brief) > 200 else brief
                print(f"\n📋 Document preview:\n{preview}\n")
                
        except FileNotFoundError as e:
            print(f"\n❌ File Not Found Error:")
            print(f"   {e}")
            print(f"\n💡 Tip: Check the file path and ensure the file exists.")
            sys.exit(1)
            
        except ImportError as e:
            print(f"\n❌ Missing Dependency:")
            print(f"   {e}")
            print(f"\n💡 Quick Fix:")
            print(f"   pip install sota-agent-framework[documents]")
            sys.exit(1)
            
        except ValueError as e:
            print(f"\n❌ Document Processing Error:")
            print(f"   {e}")
            print(f"\n💡 Possible Solutions:")
            print(f"   - Ensure the document is not corrupted")
            print(f"   - Try converting to plain text or markdown")
            print(f"   - Check if the document is password-protected")
            sys.exit(1)
            
        except RuntimeError as e:
            print(f"\n❌ Unexpected Error:")
            print(f"   {e}")
            print(f"\n💡 Please report this issue with the document type and error message.")
            sys.exit(1)
            
        except Exception as e:
            print(f"\n❌ Unknown Error:")
            print(f"   {str(e)}")
            print(f"\n💡 Try converting your document to .txt or .md format first.")
            sys.exit(1)
    elif args.interactive:
        print("\n🏗️  SOTA Agent Framework - Architecture Advisor\n")
        print("Describe your use case in natural language, and I'll recommend")
        print("the optimal architecture from our framework.\n")
        brief = input("📝 Use case brief: ").strip()
    elif args.brief:
        brief = ' '.join(args.brief)
    else:
        parser.print_help()
        sys.exit(1)
    
    if not brief:
        print("❌ Error: Please provide a use case description")
        sys.exit(1)
    
    # Show capabilities overview (skip in JSON-only mode)
    if not args.json:
        ArchitectureAdvisor.print_capabilities_overview()
    
    # Analyze and recommend
    recommendation = advisor.analyze_brief(brief)
    
    # Output
    if args.json:
        # JSON-only mode (for automation/scripting)
        import json
        output = {
            'level': recommendation.level.value,
            'level_name': recommendation.level_name,
            'confidence': recommendation.confidence,
            'input_schema': recommendation.input_schema,
            'output_schema': recommendation.output_schema,
            'features': recommendation.features,
            'integrations': recommendation.integrations,
            'reasoning': recommendation.reasoning,
            'estimated_hours': recommendation.estimated_hours,
            'generation_params': recommendation.generation_params
        }
        print(json.dumps(output, indent=2))
    else:
        # Human-readable + JSON details (default)
        advisor.print_recommendation(recommendation, show_json=True)
        
        # Interactive level selection
        selected_level = recommendation.level.value
        if args.select:
            selected_level = advisor.prompt_level_selection(recommendation)
        
        # Show next steps with selected level
        print("\n" + "="*70)
        print("🚀 Next Steps")
        print("="*70)
        print(f"\n1️⃣  Start learning at Level {selected_level}:")
        print(f"   sota-learn start {selected_level}")
        print()
        print(f"2️⃣  Generate project scaffold:")
        print(f"   sota-generate my_project --level {selected_level}")
        print()
        print(f"3️⃣  Explore interactively:")
        print(f"   sota-setup")
        print()
        
        if not args.select:
            print(f"💡 Tip: Use --select flag for interactive level selection")
            print(f"   sota-architect --file your_brief.pdf --select")
        
        print("="*70 + "\n")


if __name__ == '__main__':
    main()

